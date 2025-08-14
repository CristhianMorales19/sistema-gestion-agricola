#!/usr/bin/env python3
"""
Script para convertir casos de uso de Markdown a formato Word (.docx)
para adjuntar en Azure DevOps

Autor: Sistema de Gestión Agrícola UNA
Fecha: Diciembre 2024
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn
import re

def setup_document_styles(doc):
    """Configura los estilos del documento Word"""
    
    # Estilo para título principal
    if 'Título Principal' not in [s.name for s in doc.styles]:
        title_style = doc.styles.add_style('Título Principal', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)
    
    # Estilo para encabezados nivel 1
    if 'Encabezado 1 Custom' not in [s.name for s in doc.styles]:
        h1_style = doc.styles.add_style('Encabezado 1 Custom', WD_STYLE_TYPE.PARAGRAPH)
        h1_style.font.name = 'Arial'
        h1_style.font.size = Pt(14)
        h1_style.font.bold = True
        h1_style.paragraph_format.space_before = Pt(12)
        h1_style.paragraph_format.space_after = Pt(6)
    
    # Estilo para encabezados nivel 2
    if 'Encabezado 2 Custom' not in [s.name for s in doc.styles]:
        h2_style = doc.styles.add_style('Encabezado 2 Custom', WD_STYLE_TYPE.PARAGRAPH)
        h2_style.font.name = 'Arial'
        h2_style.font.size = Pt(12)
        h2_style.font.bold = True
        h2_style.paragraph_format.space_before = Pt(6)
        h2_style.paragraph_format.space_after = Pt(3)
    
    # Estilo para texto normal
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)

def add_header_footer(doc, caso_uso_id, nombre_caso_uso):
    """Agrega encabezado y pie de página al documento"""
    
    section = doc.sections[0]
    
    # Encabezado
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"Universidad Nacional de Costa Rica - {caso_uso_id}: {nombre_caso_uso}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Pie de página
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = "Sistema de Control y Planificación de Mano de Obra Agroindustrial"
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def parse_markdown_to_word(markdown_file_path, output_dir):
    """Convierte un archivo Markdown a Word manteniendo el formato académico"""
    
    # Leer el archivo Markdown
    with open(markdown_file_path, 'r', encoding='utf-8') as file:
        content = file.read()
    
    # Crear nuevo documento Word
    doc = Document()
    setup_document_styles(doc)
    
    # Extraer información básica del caso de uso
    caso_uso_match = re.search(r'\*\*Caso de Uso:\*\* (.+)', content)
    caso_uso_info = caso_uso_match.group(1) if caso_uso_match else "Caso de Uso"
    
    # Extraer ID
    id_match = re.search(r'- \*\*ID:\*\* (.+)', content)
    caso_uso_id = id_match.group(1) if id_match else "CU-XXX"
    
    # Extraer nombre
    nombre_match = re.search(r'- \*\*Nombre:\*\* (.+)', content)
    nombre_caso_uso = nombre_match.group(1) if nombre_match else "Nombre del Caso de Uso"
    
    # Agregar encabezado y pie de página
    add_header_footer(doc, caso_uso_id, nombre_caso_uso)
    
    # Título principal del documento
    title_para = doc.add_paragraph()
    title_para.style = 'Título Principal'
    title_para.add_run("ESPECIFICACIÓN DE CASO DE USO").bold = True
    
    # Información del sistema
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(f"""
Sistema: Sistema de Control y Planificación de Mano de Obra Agroindustrial
{caso_uso_info}
Versión: 1.0
Fecha: Diciembre 2024
Universidad Nacional de Costa Rica
Escuela de Informática
""")
    run.font.size = Pt(11)
    
    # Línea separadora
    doc.add_paragraph("=" * 80)
    
    # Procesar el contenido Markdown línea por línea
    lines = content.split('\n')
    in_code_block = False
    in_table = False
    table_headers = []
    
    for i, line in enumerate(lines):
        line = line.strip()
        
        # Saltar las primeras líneas ya procesadas
        if i < 10:  # Saltar título y metadatos iniciales
            continue
            
        # Detectar bloques de código
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            # Agregar línea de código con formato monospace
            code_para = doc.add_paragraph(line)
            code_para.style = 'Normal'
            for run in code_para.runs:
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            continue
        
        # Procesar encabezados
        if line.startswith('# '):
            para = doc.add_paragraph()
            para.style = 'Encabezado 1 Custom'
            para.add_run(line[2:]).bold = True
            
        elif line.startswith('## '):
            para = doc.add_paragraph()
            para.style = 'Encabezado 1 Custom'
            para.add_run(line[3:])
            
        elif line.startswith('### '):
            para = doc.add_paragraph()
            para.style = 'Encabezado 2 Custom'
            para.add_run(line[4:])
            
        # Procesar listas
        elif line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph()
            para.style = 'Normal'
            para.paragraph_format.left_indent = Inches(0.25)
            
            # Procesar texto con formato (negrita, cursiva)
            text = line[2:]
            add_formatted_text(para, text)
            
        elif line.startswith(('1. ', '2. ', '3. ', '4. ', '5. ', '6. ', '7. ', '8. ', '9. ')):
            para = doc.add_paragraph()
            para.style = 'Normal'
            para.paragraph_format.left_indent = Inches(0.25)
            add_formatted_text(para, line)
            
        # Procesar tablas simples
        elif '|' in line and line.count('|') >= 2:
            if not in_table:
                in_table = True
                table_headers = [cell.strip() for cell in line.split('|')[1:-1]]
                # Crear tabla
                table = doc.add_table(rows=1, cols=len(table_headers))
                table.style = 'Table Grid'
                
                # Agregar encabezados
                header_cells = table.rows[0].cells
                for j, header in enumerate(table_headers):
                    header_cells[j].text = header
                    # Formato de encabezado
                    for paragraph in header_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
            elif line.startswith('|---') or line.startswith('| ---'):
                # Línea separadora de tabla, ignorar
                continue
            else:
                # Fila de datos
                cells_data = [cell.strip() for cell in line.split('|')[1:-1]]
                if len(cells_data) == len(table_headers):
                    row_cells = table.add_row().cells
                    for j, cell_data in enumerate(cells_data):
                        row_cells[j].text = cell_data
        else:
            in_table = False
            
        # Procesar texto normal
        if line and not line.startswith(('#', '-', '*', '|')) and not line.startswith(tuple('123456789')):
            para = doc.add_paragraph()
            para.style = 'Normal'
            add_formatted_text(para, line)
    
    # Guardar el documento
    filename = Path(markdown_file_path).stem + '.docx'
    output_path = Path(output_dir) / filename
    doc.save(output_path)
    
    return output_path

def add_formatted_text(paragraph, text):
    """Agrega texto con formato (negrita, cursiva) al párrafo"""
    
    # Procesar texto en negrita **texto**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            # Texto en negrita
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            # Texto en cursiva
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            # Texto normal
            paragraph.add_run(part)

def main():
    """Función principal del script"""
    
    # Directorio de casos de uso en Markdown
    markdown_dir = Path(__file__).parent
    word_dir = markdown_dir / 'word'
    
    # Crear directorio de salida si no existe
    word_dir.mkdir(exist_ok=True)
    
    # Buscar archivos Markdown de casos de uso
    markdown_files = list(markdown_dir.glob('CU-*.md'))
    
    if not markdown_files:
        print("No se encontraron archivos de casos de uso (CU-*.md)")
        return
    
    print(f"Convirtiendo {len(markdown_files)} casos de uso a formato Word...")
    
    converted_files = []
    for md_file in markdown_files:
        try:
            output_path = parse_markdown_to_word(md_file, word_dir)
            converted_files.append(output_path)
            print(f"✓ Convertido: {md_file.name} -> {output_path.name}")
            
        except Exception as e:
            print(f"✗ Error convertiendo {md_file.name}: {str(e)}")
    
    print(f"\n¡Conversión completada!")
    print(f"Archivos Word generados en: {word_dir}")
    print(f"Total convertidos: {len(converted_files)}")
    
    # Mostrar archivos generados
    print("\nArchivos generados:")
    for file_path in converted_files:
        print(f"  - {file_path.name}")

if __name__ == "__main__":
    main()
